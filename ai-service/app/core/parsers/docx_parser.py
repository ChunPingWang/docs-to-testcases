from docx import Document
from pathlib import Path


def _table_to_markdown(table) -> str:
    """Convert a docx table to markdown format."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)
    if not rows:
        return ""
    header = rows[0]
    md = "| " + " | ".join(header) + " |\n"
    md += "| " + " | ".join("---" for _ in header) + " |\n"
    for row in rows[1:]:
        md += "| " + " | ".join(row) + " |\n"
    return md.strip()


def parse_docx(file_path: str) -> dict:
    """Parse a .docx file and extract text with structural metadata.

    Tables are converted to markdown and embedded in-place within section content
    rather than appended at the end.
    """
    doc = Document(file_path)
    sections = []
    current_heading = ""
    current_text: list[str] = []
    heading_path: list[str] = []

    # Build a mapping of table positions by checking document body elements
    # docx tables and paragraphs are interleaved in doc.element.body
    table_set = set(id(t) for t in doc.tables)
    table_map = {id(t): t for t in doc.tables}

    tables: list[str] = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            # It's a paragraph
            from docx.text.paragraph import Paragraph
            para = Paragraph(element, doc)
            style_name = para.style.name if para.style else ""

            if style_name.startswith("Heading"):
                # Save previous section
                if current_text:
                    sections.append({
                        "title": current_heading,
                        "heading_path": " > ".join(heading_path),
                        "content": "\n".join(current_text),
                    })
                    current_text = []

                level = int(style_name.replace("Heading ", "").strip() or "1")
                current_heading = para.text.strip()
                heading_path = heading_path[: level - 1] + [current_heading]
            elif para.text.strip():
                current_text.append(para.text.strip())

        elif tag == "tbl":
            # It's a table — find the matching Table object
            from docx.table import Table
            tbl = Table(element, doc)
            md = _table_to_markdown(tbl)
            if md:
                current_text.append(md)
                tables.append(md)

    # Save last section
    if current_text:
        sections.append({
            "title": current_heading,
            "heading_path": " > ".join(heading_path),
            "content": "\n".join(current_text),
        })

    full_text = "\n\n".join(s["content"] for s in sections)

    return {
        "text": full_text,
        "sections": sections,
        "tables": tables,
        "metadata": {
            "filename": Path(file_path).name,
            "section_count": len(sections),
            "table_count": len(tables),
        },
    }
